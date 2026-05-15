import json
import os
import tempfile
from pathlib import Path
from typing import Any, Final

import pandas as pd
from nicegui import ui

try:
    import docx
    from docx.shared import Inches

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from socrata_toolkit.viz import export_plotly_figure, plot_sidewalk_anatomy

from .engineering import SIDEWALK_MATERIALS, SidewalkAnatomy, SidewalkZone

CACHE_FILE: Final[Path] = Path("material_costs_cache.json")

# Default unit costs for demonstration (Dollars per sqft)
DEFAULT_MATERIAL_COSTS: Final[dict[str, float]] = {
    "Unpigmented Concrete": 15.0,
    "Pigmented Concrete (Dark)": 18.0,
    "Pigmented Concrete (Historic)": 22.0,
    "Detectable Warning Surface": 45.0,
    "PICP": 25.0,
    "Granite Slab": 55.0,
    "Asphaltic Concrete": 12.0,
}


class SidewalkSandboxBlock:
    """
    Interactive GUI block for real-time manipulation of sidewalk anatomy.
    Adjusts dimensions via drag-and-drop sliders and evaluates ADA compliance on the fly.
    """

    def __init__(self) -> None:
        # Initialize a default sidewalk segment to play with
        self.anatomy = SidewalkAnatomy(
            segment_id="Interactive-Sandbox-1",
            length_ft=20.0,
            frontage_zone=SidewalkZone("Frontage", 2.0, "Unpigmented Concrete", 1.5),
            pedestrian_zone=SidewalkZone("Clear Path", 6.0, "Unpigmented Concrete", 1.5),
            furniture_zone=SidewalkZone("Planting Strip", 4.0, "PICP", 2.0),
            curb_zone=SidewalkZone("Curb", 0.5, "Granite Slab", 0.0),
        )
        self.ada_status_container: ui.column | None = None
        self.cost_table_container: ui.column | None = None
        self.chart: ui.plotly | None = None

    def render(self) -> None:
        """Renders the drag-and-drop sandbox and schematic."""
        with ui.card().classes("w-full border border-gray-200 shadow-sm"):
            with ui.tabs().classes("w-full") as tabs:
                sandbox_tab = ui.tab("Sandbox", icon="architecture")
                config_tab = ui.tab("Configuration", icon="settings")

            with ui.tab_panels(tabs, value=sandbox_tab).classes("w-full bg-transparent p-0"):
                with ui.tab_panel(sandbox_tab).classes("p-0"):
                    with ui.row().classes("w-full items-center justify-between mb-4"):
                        _ = ui.label("📐 Sidewalk Anatomy Sandbox").classes(
                            "text-xl font-bold text-gray-800 dark:text-gray-200"
                        )

                        with ui.dropdown_button(
                            "Export Options", icon="download", auto_close=True
                        ).classes("bg-blue-600 text-white outline"):
                            ui.item("Generate IFA Memo (.docx)", on_click=self.export_ifa_memo)
                            ui.item("Export Schematic (.pdf)", on_click=self.export_pdf)
                            ui.separator()
                            ui.item("Export Data (.csv)", on_click=lambda: self.export_data("csv"))
                            ui.item(
                                "Export Data (.xlsx)", on_click=lambda: self.export_data("xlsx")
                            )
                            ui.item(
                                "Export Data (.json)", on_click=lambda: self.export_data("json")
                            )

                    with ui.row().classes("w-full gap-4"):
                        # Left Column: Interactive Controls & Compliance Output
                        with ui.column().classes(
                            "w-1/3 bg-gray-50 dark:bg-gray-800 p-4 rounded shadow-inner"
                        ):
                            _ = ui.label("Zone Widths (ft)").classes("font-bold text-lg mb-2")

                            self._create_zone_controls("Frontage Zone", self.anatomy.frontage_zone)
                            self._create_zone_controls(
                                "Pedestrian Clear Path", self.anatomy.pedestrian_zone
                            )
                            self._create_zone_controls(
                                "Furniture/Planting", self.anatomy.furniture_zone
                            )
                            self._create_zone_controls("Curb Zone", self.anatomy.curb_zone)

                            ui.separator().classes("my-4")
                            _ = ui.label("Live ADA Evaluation").classes("font-bold text-lg mb-2")
                            self.ada_status_container = ui.column().classes("w-full")
                            self._update_ada_readout()

                            ui.separator().classes("my-4")
                            _ = ui.label("Cost Estimate").classes("font-bold text-lg mb-2")
                            self.cost_table_container = ui.column().classes("w-full")
                            self._update_cost_readout()

                        # Right Column: Live Plotly Schematic
                        with ui.column().classes("w-3/5 flex-grow"):
                            # Create initial chart
                            self.chart = ui.plotly(self._generate_current_figure()).classes(
                                "w-full h-[500px]"
                            )

                with ui.tab_panel(config_tab).classes(
                    "p-4 bg-gray-50 dark:bg-gray-800 rounded shadow-inner"
                ):
                    ui.label("Global Material Unit Costs ($/sqft)").classes(
                        "text-xl font-bold mb-4"
                    )
                    ui.markdown(
                        "Adjust the baseline unit costs for each material. Changes are cached globally and reflect instantly in the Sandbox tab."
                    )

                    current_costs = self.get_costs()
                    with ui.grid(columns=2).classes("w-full gap-4 max-w-3xl mt-4"):
                        for mat in SIDEWALK_MATERIALS.keys():
                            val = current_costs.get(mat, 20.0)
                            with ui.row().classes(
                                "w-full items-center justify-between p-2 border border-gray-200 rounded bg-white dark:bg-gray-900"
                            ):
                                ui.label(mat).classes("font-semibold text-sm w-1/2")
                                ui.number(
                                    value=val,
                                    format="%.2f",
                                    step=0.5,
                                    on_change=lambda e, m=mat: self._update_single_cost(m, e.value),
                                ).classes("w-1/3")

    def _create_zone_controls(self, label_text: str, zone_obj: SidewalkZone):
        """Helper to create connected sliders and dropdowns for a given sidewalk zone."""
        with ui.column().classes("w-full mb-4"):
            ui.label(label_text).classes("text-sm font-semibold")

            with ui.row().classes("w-full items-center gap-2"):
                slider = ui.slider(
                    min=0.0,
                    max=15.0,
                    step=0.5,
                    value=zone_obj.width_ft,
                    on_change=lambda e, z=zone_obj: self._on_width_change(z, e.value),
                ).classes("flex-grow")
                ui.label().bind_text_from(slider, "value", backward=lambda v: f"{v:.1f}'").classes(
                    "text-sm font-mono w-8"
                )

            material_opts = list(SIDEWALK_MATERIALS.keys())
            ui.select(
                material_opts,
                value=zone_obj.material,
                label="Material",
                on_change=lambda e, z=zone_obj: self._on_material_change(z, e.value),
            ).classes("w-full text-xs")

    def _on_width_change(self, zone_obj: SidewalkZone, new_width: float):
        """Callback triggered when a slider is dragged."""
        # 1. Update the underlying python math model
        zone_obj.width_ft = new_width
        # 2. Push the new math to the Plotly figure without a full page reload
        self.chart.update_figure(self._generate_current_figure())
        # 3. Re-run ADA rules
        self._update_ada_readout()
        self._update_cost_readout()

    def _on_material_change(self, zone_obj: SidewalkZone, new_material: str):
        """Callback triggered when a material dropdown is changed."""
        zone_obj.material = new_material
        self.chart.update_figure(self._generate_current_figure())
        self._update_cost_readout()

    def get_costs(self):
        """Retrieves global costs from local JSON cache or fallback defaults."""
        if CACHE_FILE.exists():
            try:
                return json.loads(CACHE_FILE.read_text())
            except Exception:
                pass
        return DEFAULT_MATERIAL_COSTS.copy()

    def _update_single_cost(self, material: str, new_cost: float):
        if new_cost is None:
            return
        costs = self.get_costs()
        costs[material] = new_cost
        CACHE_FILE.write_text(json.dumps(costs, indent=2))
        self._update_cost_readout()

    def _generate_current_figure(self) -> Any:
        """Regenerates the Plotly figure based on current anatomy dimensions."""
        geo_vectors = self.anatomy.to_vector_geojson(include_corner_ramp=True)
        return plot_sidewalk_anatomy(geo_vectors, title="Live Dynamic Schematic")

    def _update_ada_readout(self):
        """Updates the text describing the real-time ADA compliance status."""
        self.ada_status_container.clear()
        with self.ada_status_container:
            status = self.anatomy.evaluate_ada_compliance()
            if status["is_compliant"]:
                ui.label("✅ Fully Compliant").classes("text-green-600 font-bold")
            else:
                ui.label("❌ Compliance Violations:").classes("text-red-600 font-bold")
                for issue in status["compliance_issues"]:
                    ui.label(f"• {issue}").classes("text-red-500 text-sm ml-2")

    def _get_cost_data(self):
        """Calculates area and estimated cost based on current zone dimensions."""
        length = self.anatomy.length_ft
        zones = [
            self.anatomy.frontage_zone,
            self.anatomy.pedestrian_zone,
            self.anatomy.furniture_zone,
            self.anatomy.curb_zone,
        ]
        data = []
        total_cost = 0.0
        current_costs = self.get_costs()

        for z in zones:
            if z.width_ft > 0:
                area = z.width_ft * length
                unit_cost = current_costs.get(z.material, 20.0)  # Default $20/sqft if not mapped
                cost = area * unit_cost
                total_cost += cost
                data.append(
                    {
                        "Zone": z.name,
                        "Material": z.material,
                        "Area (sqft)": f"{area:.1f}",
                        "Est. Cost": f"${cost:,.2f}",
                    }
                )
        return data, total_cost

    def _update_cost_readout(self):
        """Updates the interactive cost table."""
        self.cost_table_container.clear()
        data, total_cost = self._get_cost_data()

        with self.cost_table_container:
            if data:
                ui.aggrid(
                    {"columnDefs": [{"field": k} for k in data[0].keys()], "rowData": data}
                ).classes("w-full h-40")
                ui.label(f"Total Estimate: ${total_cost:,.2f}").classes(
                    "text-lg font-bold mt-2 text-right w-full text-blue-700 dark:text-blue-400"
                )
            else:
                ui.label("No zones defined.").classes("italic text-gray-500")

    def export_pdf(self):
        """Exports the current schematic to a PDF file."""
        ui.notify("Generating PDF... (Requires kaleido package)", type="info")

        fig = self._generate_current_figure()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            base_path = tmp.name[:-4]  # Remove .pdf extension for export_plotly_figure

        try:
            saved_paths = export_plotly_figure(fig, base_path, formats=["pdf"])
            if saved_paths:
                ui.download(saved_paths[0], "sidewalk_anatomy_schematic.pdf")
            else:
                ui.notify(
                    'PDF export failed. Check if "kaleido" package is installed.', type="warning"
                )
        except Exception as e:
            ui.notify(f"Export error: {e}", type="negative")

    def export_data(self, fmt: str):
        """Exports the raw dimension and cost data."""
        data, _ = self._get_cost_data()
        if not data:
            ui.notify("No data to export.", type="warning")
            return

        df = pd.DataFrame(data)
        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{fmt}") as tmp:
            if fmt == "csv":
                df.to_csv(tmp.name, index=False)
            elif fmt == "xlsx":
                df.to_excel(tmp.name, index=False)
            elif fmt == "json":
                df.to_json(tmp.name, orient="records", indent=2)
            ui.download(tmp.name, f"{self.anatomy.segment_id}_estimate.{fmt}")

    def export_ifa_memo(self):
        """Generates an official Word Document (.docx) memo appending the charts and costs."""
        if not DOCX_AVAILABLE:
            ui.notify(
                "python-docx is not installed. Run 'pip install python-docx'", type="negative"
            )
            return

        ui.notify("Generating IFA Memo... Please wait.", type="info")

        # 1. Capture the visual diagram as a temporary PNG
        fig = self._generate_current_figure()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as img_tmp:
            img_base_path = img_tmp.name[:-4]

        try:
            saved_paths = export_plotly_figure(fig, img_base_path, formats=["png"])
            img_path = saved_paths[0] if saved_paths else None
        except Exception as e:
            ui.notify(f"Diagram generation failed: {e}. Generating text-only memo.", type="warning")
            img_path = None

        # 2. Build the Official Document
        doc = docx.Document()
        doc.add_heading("NYC DOT Sidewalk IFA Justification Memo", 0)

        doc.add_paragraph(f"Segment ID: {self.anatomy.segment_id}")
        doc.add_paragraph(
            f"Total Width: {self.anatomy.total_width_ft} ft | Length: {self.anatomy.length_ft} ft"
        )

        doc.add_heading("ADA Compliance Review", level=2)
        status = self.anatomy.evaluate_ada_compliance()
        if status["is_compliant"]:
            doc.add_paragraph("STATUS: Fully Compliant with ADA Guidelines.")
        else:
            doc.add_paragraph("STATUS: Compliance Violations Detected:")
            for issue in status["compliance_issues"]:
                doc.add_paragraph(f"{issue}", style="List Bullet")

        doc.add_heading("Cost Estimate Breakdown", level=2)
        data, total = self._get_cost_data()
        df = pd.DataFrame(data)

        # Create Word Table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = col_name

        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                row_cells[i].text = str(val)

        doc.add_paragraph(f"\nTotal Estimated Cost: ${total:,.2f}").bold = True

        # 3. Attach Diagram
        if img_path and os.path.exists(img_path):
            doc.add_heading("Infrastructure Schematic", level=2)
            doc.add_picture(img_path, width=Inches(6.0))

        # 4. Save and Trigger Download
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as doc_tmp:
            doc.save(doc_tmp.name)
            ui.download(doc_tmp.name, f"IFA_Memo_{self.anatomy.segment_id}.docx")

        ui.notify("Memo successfully created and downloaded!", type="positive")
